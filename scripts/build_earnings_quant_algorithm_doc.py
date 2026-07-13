from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Gistify_Earnings_Quant_Algoritmasi_Teknik_Dokumani.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
CYAN = "168AAD"
INK = "1F2937"
MUTED = "64748B"
LIGHT = "E8EEF5"
LIGHTER = "F4F6F9"
GOLD = "9A6700"
RED = "9B1C1C"
GREEN = "166534"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="CBD5E1", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_run(run, size=11, bold=False, color=INK, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Sayfa ")
    set_run(run, 9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("GISTIFY  /  EARNINGS QUANT INTELLIGENCE")
    set_run(hr, 8.5, bold=True, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_title_page(doc):
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TEKNIK REFERANS RAPORU")
    set_run(r, 10, bold=True, color=CYAN)
    p.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Gistify Earnings Quant\nAlgoritmasi")
    set_run(r, 30, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Gunluk piyasa kosullarina uyarlanan opsiyon strateji, risk ve uyari motoru")
    set_run(r, 14, color=BLUE)
    p.paragraph_format.space_after = Pt(32)

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(callout, [7200], indent=1080)
    set_table_borders(callout, color="B9DDE8", size="8")
    set_cell_shading(callout.cell(0, 0), "EFF8FB")
    cp = callout.cell(0, 0).paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run("Kapsam: veri akisindan strateji secimine; BSM modellemesinden Kelly boyutlamaya; Greeks, margin, alarm ve UI sunumuna kadar uygulanan tum karar zinciri.")
    set_run(cr, 10.5, color=NAVY)

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Surum 1.0  |  11 Temmuz 2026  |  Teknik Dokumantasyon")
    set_run(r, 10, bold=True, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Uygulama kaynagi: Gistify earningsStrategy + earningsQuantEngine")
    set_run(r, 9, italic=True, color=MUTED)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run(r1, bold=True, color=NAVY)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run(r2, italic=italic)
    else:
        r = p.add_run(text)
        set_run(r, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_run(p.add_run(item))


def add_numbers(doc, items):
    numbering = doc.part.numbering_part.element
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1

    list_style = doc.styles["List Number"]
    base_num_id = int(list_style._element.pPr.numPr.numId.val)
    base_num = next(node for node in numbering.findall(qn("w:num")) if int(node.get(qn("w:numId"))) == base_num_id)
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_node)
        p_pr.append(num_pr)
        set_run(p.add_run(item))


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(header)
        set_run(r, font_size, bold=True, color=NAVY)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "F8FAFC")
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run(r, font_size, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, text, tone="info"):
    palette = {
        "info": ("EFF8FB", CYAN),
        "risk": ("FEF2F2", RED),
        "warning": ("FFF8E7", GOLD),
        "success": ("F0FDF4", GREEN),
    }
    fill, accent = palette[tone]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=accent, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}\n")
    set_run(r, 10.5, bold=True, color=accent)
    r = p.add_run(text)
    set_run(r, 10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def keep_table_rows(table):
    for row in table.rows:
        prevent_row_split(row)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_together = True


def build_document():
    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    add_heading(doc, "Dokumanin Amaci ve Okuma Rehberi", 1)
    add_para(doc, "Bu dokuman, Gistify Earnings Strategy sayfasina eklenen kantitatif karar motorunun uygulanan gercek davranisini aciklar. Anlatim; veri kaynaklari, normalizasyon, gostergeler, skor birlestirme, strateji secimi, opsiyon bacaklarinin modellenmesi, risk metrikleri, alarm kurallari, yenileme servisi ve kullanici arayuzu boyunca ilerler.")
    add_callout(doc, "Onemli kapsam siniri", "Sistem karar destegi ve modelleme motorudur. Canli opsiyon zinciri bagli olmadiginda BSM ile uretilen primler, Greeks ve olasiliklar gercek bid/ask kotasyonu degildir. Arayuz bu durumu MODEL / VERIFY etiketiyle acikca gosterir.", "warning")

    add_heading(doc, "Icindekiler", 2)
    toc_items = [
        "1. Yonetici ozeti ve sistem hedefi",
        "2. Uctan uca mimari",
        "3. Veri kaynaklari ve kalite katmanlari",
        "4. Rapor parser ve geriye uyumluluk",
        "5. Gunluk momentum motoru",
        "6. Opsiyon akis ve volatilite modeli",
        "7. Bilesik skor ve yon karari",
        "8. Dinamik strateji secim agaci",
        "9. Strike, bacak ve prim modelleme",
        "10. Payoff, POP, EV, Kelly, Greeks ve margin",
        "11. Uyari ve pozisyon yonetim motoru",
        "12. Yenileme, API ve arayuz akisi",
        "13. Testler, sinirlar ve gelisim yol haritasi",
        "Ek A. Formul referansi",
        "Ek B. Pseudocode",
        "Ek C. Terimler sozlugu",
    ]
    add_bullets(doc, toc_items)
    doc.add_page_break()

    add_heading(doc, "1. Yonetici Ozeti ve Sistem Hedefi", 1)
    add_para(doc, "Motorun temel amaci, aylik earnings raporunda yazan statik stratejiyi tek dogru kabul etmek yerine her yenilemede guncel fiyat davranisi, momentum, Call/Put dengesi, IV rejimi ve earnings'e kalan gunu yeniden tartmaktir. Kosullar degismisse sistem stratejiyi degistirir, degisimin nedenini puan bazinda aciklar ve yeni yapinin riskini hesaplar.")
    add_heading(doc, "1.1 Uretilen ana ciktilar", 2)
    add_bullets(doc, [
        "Her ticker icin bullish, neutral veya bearish yon sinifi.",
        "-100 ile +100 arasinda bilesik karar skoru ve yuzde guven degeri.",
        "Gunluk degisim, 5/20 gunluk getiri, RSI(14), gerceklesen volatilite ve momentum etiketi.",
        "Call/Put orani, akisin call agirlikli, dengeli veya put agirlikli olduguna dair yorum.",
        "IV Rank ve DTE'ye uygun Iron Condor, Calendar Spread, Bull/Bear Credit Spread veya Debit Spread secimi.",
        "Model strike'lari, bacaklar, BSM proxy primleri, beklenen hareket, POP ve breakeven seviyeleri.",
        "Maksimum kar/zarar, return-on-risk, Reg-T margin ve ceyrek-Kelly pozisyon boyutu.",
        "Gunluk fiyat soku, RSI asiriligi, earnings gamma penceresi ve asiri IV uyarilari.",
    ])
    add_heading(doc, "1.2 Tasarim ilkeleri", 2)
    add_table(doc, ["Ilke", "Uygulama"], [
        ["Aciklanabilirlik", "Her secim; fiyat, akis, RSI ve IV/DTE katkilarini metin olarak gosterir."],
        ["Tanımlı risk", "Yonlu goruslerde mumkun oldugunca spread; notr goruste Iron Condor kullanilir."],
        ["Dogruluk etiketi", "Live veri ile model/proxy veri birbirinden ayrilir."],
        ["Geriye uyumluluk", "Eski ve yeni markdown rapor bicimleri ayni servis tarafindan okunur."],
        ["Graceful degradation", "FMP veya zincir verisi yoksa sayfa cokmez; rapor girdileriyle model moduna gecer."],
        ["Islem disiplini", "Limit emir, likidite dogrulamasi, kar alma ve zaman stopu karar metnine dahildir."],
    ], [2100, 7260])

    doc.add_page_break()
    add_heading(doc, "2. Uctan Uca Mimari", 1)
    add_para(doc, "Sistem bes katmandan olusur. Katmanlar arasindaki ayrim, veri hatasinin karar mantigina sessizce sizmasini engeller.")
    add_table(doc, ["Katman", "Sorumluluk", "Ana bilesen"], [
        ["1. Kaynak", "Master markdown, FMP fiyat gecmisi, ortam ayarlari", "earningreport/, FMP_API_KEY"],
        ["2. Normalizasyon", "Takvim, strateji, CPR, IV Rank ve fiyat alanlarini ortak tipe donusturur", "earningsStrategy.ts"],
        ["3. Quant motor", "Momentum, volatilite, skor, strateji, payoff ve uyarilari hesaplar", "earningsQuantEngine.ts"],
        ["4. API/senkronizasyon", "Snapshot'i periyodik yeniler ve istemciye sunar", "/api/earnings/strategy"],
        ["5. Sunum", "Command Center, ticker secimi, kartlar ve veri kalitesi etiketleri", "EarningsQuantCommandCenter.tsx"],
    ], [1800, 4200, 3360])
    add_heading(doc, "2.1 Calisma akisi", 2)
    add_numbers(doc, [
        "Servis, earningreport klasorundeki en yeni *_Earnings_Opsiyon_Master_Stratejisi.md dosyasini bulur.",
        "Markdown parser raporu ortak EarningsStrategyData modeline donusturur.",
        "Ilk 12 oncelikli ticker icin FMP'den yaklasik 70 takvim gunluk EOD fiyat serisi istenir.",
        "Her strateji buildStrategyIntelligence fonksiyonundan gecirilir.",
        "Quant overview; canli kapsama, yon dagilimi, strateji degisimi ve kritik alarm sayilarini toplar.",
        "Snapshot API ile istemciye gider; sayfa command center ve strateji kartlarini yeniler.",
    ])
    add_callout(doc, "Yenileme semantigi", "Rapor dosyasi degismese bile quant katmani yeniden hesaplanir. Boylece statik rapor uzerinde gunluk fiyat kosullari degistikce karar da degisebilir.", "success")

    add_heading(doc, "3. Veri Kaynaklari ve Kalite Katmanlari", 1)
    add_heading(doc, "3.1 Master earnings raporu", 2)
    add_para(doc, "Rapor; ticker, rapor fiyati, IV Rank, CPR, earnings tarihi, oturum, ana strateji, giris/cikis, max hold ve maliyet gibi alanlari saglar. Bu alanlar her ticker icin baseline tezdir. Quant motor baseline'i test eder; uygun bulursa korur, degilse degistirir.")
    add_heading(doc, "3.2 FMP gunluk fiyat verisi", 2)
    add_para(doc, "FMP_API_KEY tanimliysa motor ilk 12 ticker icin historical-price-eod/full uc noktasini kullanir. Yaklasik 70 gunluk pencere; son fiyat, onceki kapanis, 5 ve 20 gunluk getiri, RSI ve 20 gunluk gerceklesen volatilite icin yeterli gozlem saglar. Her ticker istegi bagimsiz hata yakalama ile calisir; tek bir ticker hatasi tum snapshot'i bozmaz.")
    add_heading(doc, "3.3 Veri kalite durumlari", 2)
    add_table(doc, ["Etiket", "Anlam", "Arayuz davranisi"], [
        ["LIVE", "FMP fiyat serisi mevcut; CPR ve IV rapordan gelir", "LIVE + MODEL rozeti"],
        ["MIXED", "Canli ve rapor kaynaklari birlestirilmis ara durum", "Kaynak notu ile gosterim"],
        ["REPORT", "Fiyat gecmisi yok; rapor fiyati ve proxy model kullanilir", "MODEL / VERIFY rozeti ve bilgi alarmi"],
    ], [1500, 4700, 3160])
    add_callout(doc, "Guvenlik", "API anahtari yalnizca .env.local icindeki FMP_API_KEY degiskeninden okunur. Anahtar dokumana, API cevabina veya loglara yazilmaz. .env.local Git tarafindan dislanir.", "info")

    add_heading(doc, "4. Rapor Parser ve Geriye Uyumluluk", 1)
    add_para(doc, "Parser iki rapor ailesini destekler: structured vNext bicimi ve legacy tablo/bolum bicimi. Yeni Temmuz-Agustos formatinda ticker basliklari numarasizdir ve IV Rank ile CPR ayni baslikta pipe karakterleriyle ayrilir. Parser bu farki regex ve tablo-baslik eslestirmesiyle normalize eder.")
    add_heading(doc, "4.1 En yeni format icin duzeltilen noktalar", 2)
    add_bullets(doc, [
        "Numarasiz baslik: ### NFLX - $74.22 | IV Rank: 72/100 | CPR: 0.85 | NEUTRAL.",
        "Birden fazla aylik takvim tablosu tek listeye birlestirilir ve tarih+ticker ile tekillestirilir.",
        "Takvim tablosu yalnizca Tarih ve Hisse/Ticker basliklarini birlikte tasiyorsa secilir; makro risk tablosu yanlislikla takvim sayilmaz.",
        "Entry, Exit, Max Hold, Kar Hedefi ve IV Crush alanlari parametre tablosundan okunur.",
        "Eski numarali baslik ve structured report desteği korunur.",
    ])
    add_heading(doc, "4.2 Normalizasyon ornegi", 2)
    add_table(doc, ["Markdown alani", "Ortak veri alani", "Kullanim"], [
        ["IV Rank: 72/100", "strategy.ivRank = 72", "Volatilite rejimi ve strateji tipi"],
        ["CPR: 0.85", "strategy.cpr = 0.85", "Call/Put akis puani"],
        ["2026-07-16 AMC", "calendar.date/time", "DTE ve gamma alarmi"],
        ["Strateji 1: Iron Condor", "strategy.type", "Onceki strateji ve degisim kontrolu"],
        ["$74.22", "reportSpot", "Canli fiyat yoksa fallback spot"],
    ], [2400, 2600, 4360])

    add_heading(doc, "5. Gunluk Momentum Motoru", 1)
    add_heading(doc, "5.1 Gunluk fiyat degisimi", 2)
    add_para(doc, "change1d = ((Spot / PreviousClose) - 1) x 100")
    add_para(doc, "Bu metrik hem arayuzde gosterilir hem fiyat skoru icinde 5 katsayisiyla kullanilir. Mutlak gunluk hareket %4 veya daha buyukse kritik fiyat soku alarmi uretilir.")
    add_heading(doc, "5.2 Coklu pencere getirileri", 2)
    add_para(doc, "returnNd = ((Close_t / Close_(t-N)) - 1) x 100. Motor N=5 ve N=20 kullanir. Bes gunluk momentum kisa vadeli hizlanmayi, 20 gunluk momentum orta vadeli trendi temsil eder.")
    add_heading(doc, "5.3 RSI(14)", 2)
    add_para(doc, "Son 14 fiyat farkindaki pozitif degisimler kazanc, negatif degisimlerin mutlak degeri kayip olarak toplanir. RS = toplam kazanc / toplam kayip; RSI = 100 - 100/(1+RS). Kayip sifirsa RSI 100 kabul edilir.")
    add_table(doc, ["RSI bolgesi", "Yorum", "Motor tepkisi"], [
        ["RSI <= 28", "Asiri satim", "Momentum asiriligi uyarisi; tanimli risk"],
        ["28 < RSI < 50", "Zayif/negatif", "Skora negatif katki"],
        ["50 - 72", "Pozitif/normal", "Skora pozitif katki"],
        ["RSI >= 72", "Asiri alim", "Geri donus ve gap riski uyarisi"],
    ], [1700, 3000, 4660])
    add_heading(doc, "5.4 Gerceklesen volatilite", 2)
    add_para(doc, "Gunluk log getirilerinin ornek standart sapmasi hesaplanir ve sqrt(252) ile yilliklastirilir: RV20 = stdev(ln(C_t/C_(t-1))) x sqrt(252) x 100. En fazla son 21 kapanis kullanilir.")
    add_heading(doc, "5.5 Fiyat ve RSI skoru", 2)
    add_para(doc, "priceScore = clamp(change1d x 5 + return5d x 2 + return20d x 0.65, -45, +45)")
    add_para(doc, "rsiScore = clamp((RSI14 - 50) x 0.65, -18, +18)")
    add_para(doc, "Momentum skoru, arayuz icin priceScore + rsiScore olarak raporlanir. Bu skor tek basina strateji secmez; opsiyon akisi ile birlestirilir.")

    add_heading(doc, "6. Opsiyon Akis ve Volatilite Modeli", 1)
    add_heading(doc, "6.1 Call/Put Ratio (CPR)", 2)
    add_para(doc, "CPR = Put aktivitesi / Call aktivitesi olarak yorumlanir. Dolayisiyla 1'in alti call agirligini, 1'in ustu put agirligini ifade eder. Arayuzde onceki ters renk yorumu duzeltilmistir.")
    add_table(doc, ["CPR araligi", "Flow etiketi", "Skor katkisi"], [
        ["CPR < 0.80", "Guclu call agirlikli", "+22"],
        ["0.80 <= CPR < 0.95", "Call agirlikli", "+10"],
        ["0.95 <= CPR <= 1.05", "Dengeli", "0"],
        ["1.05 < CPR <= 1.25", "Put agirlikli", "-10"],
        ["CPR > 1.25", "Guclu put agirlikli", "-22"],
    ], [2200, 4300, 2860])
    add_heading(doc, "6.2 IV Rank normalizasyonu", 2)
    add_para(doc, "IV Rank rapor metninden ilk sayisal deger olarak okunur ve 0-100 arasina sikistirilir. Veri yoksa 50 varsayilir. Bu varsayim, strateji motorunun notr bir volatilite rejiminden baslamasini saglar.")
    add_heading(doc, "6.3 Modeled IV", 2)
    add_para(doc, "Canli gerceklesen volatilite varsa: modeledIV = clamp(RV20 x (0.85 + IVRank/180), 15, 120). Fiyat serisi yoksa proxy: modeledIV = clamp(18 + IVRank x 0.62, 15, 120).")
    add_callout(doc, "Proxy uyarisi", "IV Rank tek basina zımni volatilitenin seviyesi degildir. Bu donusum sadece zincir verisi olmadiginda strike ve risk senaryosu uretmek icin kullanilan kontrollu bir proxy'dir.", "warning")
    add_heading(doc, "6.4 Earnings ve islem DTE'si", 2)
    add_para(doc, "Event DTE, earnings tarihi ile simdiki zaman arasindaki yukari yuvarlanmis gun farkidir ve minimum 1 kabul edilir. Modelin fiyatlama ufku 45 gunle sinirlanir: tradeDTE = min(max(eventDTE, 1), 45). Bu sinir, uzak tarihli eventlerde beklenen hareketin anlamsiz buyumesini engeller.")
    add_heading(doc, "6.5 Beklenen hareket", 2)
    add_para(doc, "ExpectedMove$ = Spot x (modeledIV/100) x sqrt(tradeDTE/365)")
    add_para(doc, "ExpectedMove% = ExpectedMove$ / Spot x 100. Bu deger yon tahmini degil; secilen vade icin bir standart sapma benzeri fiyat araligidir.")

    add_heading(doc, "7. Bilesik Skor ve Yon Karari", 1)
    add_heading(doc, "7.1 Skor formulu", 2)
    add_para(doc, "compositeScore = round(clamp(priceScore + flowScore + rsiScore, -100, +100))")
    add_table(doc, ["Skor", "Yon", "Anlam"], [
        ["score >= +18", "Bullish", "Fiyat/RSI/akis kombinasyonu pozitif"],
        ["-18 < score < +18", "Neutral", "Sinyaller dengeli veya zayif"],
        ["score <= -18", "Bearish", "Fiyat/RSI/akis kombinasyonu negatif"],
    ], [1900, 2000, 5460])
    add_heading(doc, "7.2 Makro VIX etkisi", 2)
    add_para(doc, "VIX yon skorunu degistirmez; guven degerini azaltir. VIX >= 24 ise 4 puan, VIX >= 30 ise 8 puan guven cezasi uygulanir. Bu ayrim, makro stresin yonu degil pozisyon alma konforunu azalttigi varsayimina dayanir.")
    add_heading(doc, "7.3 Guven skoru", 2)
    add_para(doc, "confidence = clamp(52 + abs(compositeScore) x 0.35 + liveBonus - macroPenalty, 35, 91)")
    add_bullets(doc, [
        "liveBonus = 12; yalnizca FMP fiyat serisi varsa uygulanir.",
        "macroPenalty = VIX esiklerine gore 0, 4 veya 8.",
        "Guven skoru basari garantisi degil; sinyal belirginligi ve veri yeterliligi gostergesidir.",
    ])

    add_heading(doc, "8. Dinamik Strateji Secim Agaci", 1)
    add_para(doc, "Strateji secimi once earnings'e kalan gun ve IV rejimini, sonra yon sinifini kullanir. Calendar Spread kosulu diger yon kurallarindan once calisir.")
    add_table(doc, ["Kosul", "Secilen strateji", "Temel mantik"], [
        ["7 <= DTE <= 20 ve IV Rank < 55", "Calendar Spread", "IV rush ve vade farki; uzun vade Vega avantaji"],
        ["Neutral ve IV Rank >= 45", "Iron Condor", "Theta/IV primi; iki tarafli tanimli risk"],
        ["Neutral ve IV Rank < 45", "Long Straddle", "Ucuz volatilite ve buyuk hareket beklentisi"],
        ["Bullish ve IV Rank >= 60", "Bull Put Spread", "Pahali put primini tanimli riskle satma"],
        ["Bullish ve IV Rank < 60", "Bull Call Spread", "Yonlu yukselise sinirli debit maruziyeti"],
        ["Bearish ve IV Rank >= 60", "Bear Call Spread", "Pahali call primini tanimli riskle satma"],
        ["Bearish ve IV Rank < 60", "Bear Put Spread", "Yonlu dususe sinirli debit maruziyeti"],
    ], [3000, 2500, 3860])
    add_heading(doc, "8.1 Strateji degisimi", 2)
    add_para(doc, "changed = reportStrategy mevcutsa ve selectedStrategy ile ayni degilse true olur. Arayuz onceki ve yeni stratejiyi okla gosterir. Degisiklik, aciklama listesine otomatik bir gerekce olarak eklenir.")
    add_callout(doc, "Ornek", "Rapor Iron Condor onerirken CPR 0.70, pozitif 5/20 gun momentum ve yuksek IV nedeniyle skor +18'i asarsa motor Bull Put Spread'e gecebilir. Bu degisim, yon ile premium rejimini ayni yapida birlestirir.", "info")

    add_heading(doc, "9. Strike, Bacak ve Prim Modelleme", 1)
    add_heading(doc, "9.1 Strike grid", 2)
    add_para(doc, "Strike adimi spot fiyata gore secilir: Spot < $50 icin $1; $50-$150 icin $2.5; Spot >= $150 icin $5. Her ham strike en yakin grid seviyesine yuvarlanir.")
    add_heading(doc, "9.2 Expected-move tabanli seviyeler", 2)
    add_para(doc, "LowerStrike = grid(Spot - ExpectedMove$); UpperStrike = grid(Spot + ExpectedMove$); WingWidth = max(gridStep, grid(ExpectedMove$ x 0.35)).")
    add_heading(doc, "9.3 Strateji bacaklari", 2)
    add_table(doc, ["Strateji", "Model bacaklari"], [
        ["Iron Condor", "Buy lower-wing Put; Sell lower Put; Sell upper Call; Buy upper-wing Call"],
        ["Bull Put Spread", "Buy lower-wing Put; Sell lower Put"],
        ["Bear Call Spread", "Sell upper Call; Buy upper-wing Call"],
        ["Bull Call Spread", "Buy ATM Call; Sell upper Call"],
        ["Bear Put Spread", "Buy ATM Put; Sell lower Put"],
        ["Long Straddle", "Buy ATM Call + Buy ATM Put"],
        ["Calendar Spread", "Sell 7-14 DTE ATM Call; Buy en az 30 DTE ATM Call"],
    ], [2500, 6860])
    add_heading(doc, "9.4 Black-Scholes-Merton proxy", 2)
    add_para(doc, "Call = S N(d1) - K exp(-rT) N(d2)")
    add_para(doc, "Put = K exp(-rT) N(-d2) - S N(-d1)")
    add_para(doc, "d1 = [ln(S/K) + (r + sigma^2/2)T] / (sigma sqrt(T)); d2 = d1 - sigma sqrt(T).")
    add_bullets(doc, [
        "Risksiz faiz r = %4 sabit proxy olarak kullanilir.",
        "T = max(DTE,1)/365; sigma = modeledIV/100.",
        "Temettu verimi mevcut implementasyonda sifir kabul edilir.",
        "Normal CDF, deterministik erf yaklasimi ile hesaplanir.",
        "Minimum opsiyon primi $0.01 tabanina sahiptir.",
    ])

    add_heading(doc, "10. Payoff, POP, EV, Kelly, Greeks ve Margin", 1)
    add_heading(doc, "10.1 Payoff taramasi", 2)
    add_para(doc, "Motor spot fiyatinin %35'i ile %165'i arasinda 401 esit fiyat noktasi olusturur. Her noktada bacaklarin vade sonu intrinsic degeri eksi modellenen prim hesaplanir; BUY bacak +1, SELL bacak -1 isaretlidir ve kontrat carpani 100'dur.")
    add_para(doc, "MaxProfit = max(PnL grid); MaxLoss = abs(min(PnL grid)). Isaret degisimlerinin oldugu ilk iki fiyat noktasi breakeven olarak raporlanir.")
    add_heading(doc, "10.2 Probability of Profit", 2)
    add_para(doc, "POP, gercek opsiyon zinciri dagilimi olmadiginda rejimsel bir tahmindir:")
    add_bullets(doc, [
        "Iron Condor: 68 + (IVRank - 50) x 0.08.",
        "Diger kredi stratejileri: 70 - abs(compositeScore) x 0.08.",
        "Debit stratejileri: 48 + abs(compositeScore) x 0.18.",
        "Sonuc %35 ile %84 arasina sikistirilir.",
    ])
    add_heading(doc, "10.3 Beklenen deger ve ceyrek-Kelly", 2)
    add_para(doc, "Yonetilen ortalama kazanc, max karın %50'si kabul edilir: AvgWin = MaxProfit x 0.50. Expectancy = POP x AvgWin - (1-POP) x MaxLoss.")
    add_para(doc, "FullKelly = p - (1-p)/(AvgWin/MaxLoss). Uygulanan Kelly = clamp(FullKelly x 0.25, 0, 0.05) x 100. Beklenen deger negatifse Kelly sifirlanir.")
    add_callout(doc, "Neden ceyrek-Kelly?", "Tam Kelly model hatasina ve rejim degisimine cok duyarlidir. Motor sermaye korunumu icin Kelly'nin yalnizca dortte birini kullanir ve islem basina maksimum %5 ile sinirlar.", "success")
    add_heading(doc, "10.4 Greeks proxy", 2)
    add_table(doc, ["Greek", "Uygulanan yaklasim", "Yorum"], [
        ["Delta", "compositeScore / 250", "Yon skoru -0.40 ile +0.40 civarina tasinir"],
        ["Gamma", "(1/max(DTE,2)) x +0.12 long / -0.06 short", "Vade yaklastikca mutlak gamma artar"],
        ["Theta", "sign(credit/debit) x abs(netDebit)/DTE", "Kredi yapisi pozitif, debit yapisi negatif"],
        ["Vega", "sign(strategy) x Spot x sqrt(DTE/365) x 0.01", "Calendar/Long pozitif; short premium negatif"],
    ], [1300, 4650, 3410])
    add_heading(doc, "10.5 Reg-T margin", 2)
    add_para(doc, "Debit yapida RegTMargin = MaxLoss. Kredi yapida margin = max(MaxLoss, spreadWidth x 100 - netCredit x 100). Bu, tanimli-risk spread icin kanat genisligi eksi kredi mantigini korur.")

    add_heading(doc, "11. Uyari ve Pozisyon Yonetim Motoru", 1)
    add_table(doc, ["Tetik", "Seviye", "Uyari", "Onerilen aksiyon"], [
        ["abs(change1d) >= %4", "Critical", "Gunluk fiyat soku", "Zinciri yeniden fiyatla; market emir yok; boyutu yariya indir"],
        ["RSI >= 72 veya <= 28", "Warning", "Momentum asiriligi", "Ciplak yon yerine tanimli-risk spread"],
        ["DTE <= 3", "Critical", "Earnings gamma penceresi", "Yeni giris yok; risk azalt; limit emir"],
        ["IV Rank >= 80", "Warning", "Asiri pahali volatilite", "Credit yapisi ve %50 kar alma"],
        ["Canli fiyat yok", "Info", "Veri dogrulamasi gerekli", "Broker zinciri, bid/ask, OI ve gercek IV ile kontrol"],
    ], [1800, 1300, 2500, 3760], font_size=8.6)
    add_heading(doc, "11.1 Giris kurali", 2)
    add_para(doc, "DTE <= 3 ise: yeni giris yok; yalnizca risk azalt veya limit emir. Diger durumlarda bid/ask <= %5, yeterli open interest ve model fiyatina yakin limit emir ile kademeli giris onerilir.")
    add_heading(doc, "11.2 Cikis kurali", 2)
    add_bullets(doc, [
        "Kredi yapilari: maksimum karin %50'sinde kapat; 21 DTE'de veya 2x kredi zararinda cik/roll.",
        "Debit yapilari: earnings IV peak oncesi kar al; premiumda %35 zarar veya tezin bozulmasinda kapat.",
        "Earnings'e 3 gun veya daha az kaldiysa gamma/IV crush alarmi karar metninin onune gecer.",
    ])

    add_heading(doc, "12. Yenileme, API ve Arayuz Akisi", 1)
    add_heading(doc, "12.1 Sunucu senkronizasyonu", 2)
    add_para(doc, "Varsayilan poll araligi 5 dakikadir; minimum 30 saniyedir. EARNINGS_STRATEGY_POLL_INTERVAL_MS ile degistirilebilir. Her yenilemede kaynak dosyanin mtime degeri kontrol edilir. Dosya degismisse yeniden parse edilir; degismemisse mevcut normalized snapshot uzerinde quant katmani yeniden calisir.")
    add_heading(doc, "12.2 API cevabi", 2)
    add_para(doc, "/api/earnings/strategy; data, pipeline ve success alanlarini dondurur. Her strategy kaydi intelligence nesnesi tasir. data.quantOverview; canli kapsama, bullish/neutral/bearish sayilari, strateji degisimleri ve kritik alarm toplamini icerir.")
    add_heading(doc, "12.3 Command Center", 2)
    add_bullets(doc, [
        "Ust ozet: canli kapsama, yon dagilimi, degisen strateji ve kritik alarm sayisi.",
        "Ticker secici: intelligence uretilebilen tum stratejiler arasinda hizli gecis.",
        "Karar paneli: spot, 1D, RSI, CPR, IV Rank, DTE, expected move, POP, max risk ve Kelly.",
        "Bacak paneli: BUY/SELL, miktar, strike, opsiyon tipi, DTE ve modellenen prim.",
        "Aciklama: skor bilesenleri, akis yorumu, IV/DTE gerekcesi ve strateji degisimi.",
        "Alarm paneli: en onemli kritik/uyari/bilgi mesajlari ve uygulanabilir aksiyon.",
    ])
    add_heading(doc, "12.4 Strateji kartlari", 2)
    add_para(doc, "Her kart, rapor stratejisinin yanina quant motorunun guncel secimini ekler. Degisim varsa STRATEJI DEGISTI etiketi gorunur. Mini metrikler 1D, RSI, POP ve composite skoru; alt satir flow sinyalini ve guveni gosterir.")

    add_heading(doc, "13. Testler, Sinirlar ve Gelisim Yol Haritasi", 1)
    add_heading(doc, "13.1 Otomatik test kapsami", 2)
    add_table(doc, ["Test", "Dogrulanan davranis"], [
        ["Pozitif momentum + CPR 0.70 + yuksek IV", "Iron Condor'dan Bull Put Spread'e gecis"],
        ["Negatif momentum + CPR 1.35", "Put agirlikli akis ve Bear Call Spread"],
        ["Quant overview", "%100 live coverage, yon toplami ve degisim sayisi"],
        ["Yeni parser formati", "Numarasiz ticker basligi, iki aylik takvim ve Entry/Exit"],
        ["Typecheck/build", "TypeScript ve production bundle basarili"],
    ], [3300, 6060])
    add_heading(doc, "13.2 Mevcut sinirlar", 2)
    add_bullets(doc, [
        "Canli option chain, bid/ask, open interest, hacim ve strike bazli gercek IV bagli degildir.",
        "Modeled IV, IV Rank ve/veya gerceklesen volatiliteden uretilen proxy'dir.",
        "POP Monte Carlo veya risk-neutral density degil; rejimsel kontrollu tahmindir.",
        "Greeks pozisyon-level analitik proxy'dir; broker Greeks yerine gecmez.",
        "Payoff taramasi %35-%165 spot araligiyla sinirlidir; ratio spread gibi sinirsiz risk yapilari secim agacinda kullanilmaz.",
        "FMP canli kapsam ilk 12 oncelikli ticker ile sinirlidir; digerleri rapor modunda kalabilir.",
        "Makro korelasyon, beta-weighted delta ve portfolio margin stres matrisi henuz portfoy seviyesinde uygulanmamistir.",
    ])
    add_callout(doc, "Islem oncesi zorunlu dogrulama", "Broker zincirinde expiration, strike, bid/ask, OI, hacim, gercek IV, temettu ve erken kullanim riski kontrol edilmeden model bacaklari emir olarak kullanilmamalidir.", "risk")
    add_heading(doc, "13.3 Onerilen gelisim sirasi", 2)
    add_numbers(doc, [
        "Canli option-chain saglayicisi ekleyerek strike bazli IV, bid/ask, OI, volume ve Greeks al.",
        "BSM proxy yerine gercek mid fiyat + likidite filtresi ve slippage modeli kullan.",
        "POP icin Monte Carlo / risk-neutral density ve earnings gap dagilimini birlestir.",
        "Strateji sonuc defteri olustur; tahmin, giris, cikis ve gerceklesen PnL ile kalibrasyon yap.",
        "Portfoy seviyesinde beta-weighted delta, sektor korelasyonu ve -%15/+%15 stres matrisi ekle.",
        "Gercek alert delivery icin e-posta/Slack/push baglantisi kur; alarm deduplication ve cooldown uygula.",
        "Walk-forward backtest ile skor agirliklarini ve IV/DTE esiklerini rejim bazinda optimize et.",
    ])

    add_heading(doc, "Ek A. Formul Referansi", 1)
    formula_rows = [
        ["1D degisim", "((S / PrevClose) - 1) x 100"],
        ["N-gun getiri", "((Close_t / Close_(t-N)) - 1) x 100"],
        ["RSI", "100 - 100/(1 + Gain/Loss)"],
        ["RV20", "stdev(log return) x sqrt(252) x 100"],
        ["Price score", "clamp(1D x 5 + R5 x 2 + R20 x 0.65, -45, 45)"],
        ["RSI score", "clamp((RSI - 50) x 0.65, -18, 18)"],
        ["Composite", "clamp(priceScore + flowScore + rsiScore, -100, 100)"],
        ["Modeled IV live", "clamp(RV20 x (0.85 + IVRank/180), 15, 120)"],
        ["Modeled IV fallback", "clamp(18 + IVRank x 0.62, 15, 120)"],
        ["Expected move", "S x sigma x sqrt(DTE/365)"],
        ["BSM d1", "[ln(S/K) + (r + sigma^2/2)T] / (sigma sqrt(T))"],
        ["BSM d2", "d1 - sigma sqrt(T)"],
        ["Expected value", "p x AvgWin - (1-p) x MaxLoss"],
        ["Quarter Kelly", "clamp(0.25 x [p - (1-p)/b], 0, 0.05)"],
        ["Return on risk", "MaxProfit / MaxLoss x 100"],
        ["Credit Reg-T", "max(MaxLoss, Width x 100 - Credit x 100)"],
    ]
    add_table(doc, ["Metrik", "Formul"], formula_rows, [2700, 6660], font_size=9)

    add_heading(doc, "Ek B. Karar Motoru Pseudocode", 1)
    pseudo = [
        "report = parseLatestMasterMarkdown()",
        "market = loadFmpHistory(first12Tickers) or empty",
        "for strategy in report.strategies:",
        "    spot = market.price or report.price",
        "    compute change1d, return5d, return20d, RSI14, RV20",
        "    flowScore = mapCallPutRatio(strategy.cpr)",
        "    composite = clamp(priceScore + flowScore + rsiScore)",
        "    bias = bullish if >=18; bearish if <=-18; else neutral",
        "    modeledIV = liveVolAdjustedIVRank or IVRankProxy",
        "    expectedMove = spot * modeledIV * sqrt(min(DTE,45)/365)",
        "    selected = chooseStrategy(bias, IVRank, eventDTE)",
        "    legs = buildExpectedMoveLegs(selected)",
        "    premiums = priceLegsWithBSM(legs)",
        "    payoff = scanExpiryGrid(35%..165%, 401 points)",
        "    pop, expectancy, quarterKelly, greeks, margin = riskMetrics()",
        "    alerts = evaluateShockRsiGammaIvAndDataQuality()",
        "    attach intelligence to strategy",
        "overview = aggregateAllIntelligence()",
        "publish /api/earnings/strategy",
    ]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="94A3B8", size="6")
    set_cell_shading(table.cell(0, 0), "0F172A")
    p = table.cell(0, 0).paragraphs[0]
    for i, line in enumerate(pseudo):
        if i:
            p.add_run("\n")
        r = p.add_run(line)
        set_run(r, 9, color="E2E8F0", font="Consolas")

    doc.add_page_break()
    add_heading(doc, "\u00a0Ek C. Terimler Sozlugu", 1)
    glossary = [
        ["BSM", "Black-Scholes-Merton opsiyon fiyatlama modeli"],
        ["CPR", "Call/Put Ratio; bu sistemde 1 alti call, 1 ustu put agirligi"],
        ["DTE", "Vadeye veya event ufkuna kalan gun"],
        ["Expected Move", "IV ve zamanla modellenen yon bagimsiz fiyat araligi"],
        ["IV Rank", "Mevcut IV'nin tarihsel aralik icindeki yuzdelik konumu"],
        ["POP", "Probability of Profit; model karlilik olasiligi"],
        ["Kelly", "Beklenen deger ve odds ile sermaye payi formulu"],
        ["Theta", "Zaman gecisinin opsiyon degerine etkisi"],
        ["Vega", "IV degisiminin opsiyon degerine etkisi"],
        ["Gamma", "Delta'nin spot degisimine duyarliligi"],
        ["Reg-T", "Perakende hesaplar icin standart teminat yaklasimi"],
        ["IV Crush", "Earnings sonrasi zımni volatilitenin hizli gerilemesi"],
        ["Defined Risk", "Maksimum zarari bacak yapisiyla sinirli strateji"],
        ["Proxy", "Canli veri yokken kontrollu varsayimla uretilen model degeri"],
    ]
    add_table(doc, ["Terim", "Tanim"], glossary, [2200, 7160])

    add_heading(doc, "Sonuc", 1)
    add_para(doc, "Gistify Earnings Quant motoru, rapor tabanli statik fikirleri gunluk piyasa kosullariyla yeniden puanlayan, gerektiğinde strateji degistiren ve secimini risk metrikleriyle aciklayan bir karar destek katmanidir. Sistem canli fiyat verisiyle otomatiklesen; opsiyon zinciri yoklugunu ise acikca etiketleyen kontrollu bir mimariye sahiptir. Bir sonraki olgunluk asamasi, gercek option-chain ve gerceklesen sonuc defteriyle modelin kalibrasyonunu kapali donguye almaktir.")

    for table in doc.tables:
        keep_table_rows(table)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    build_document()
